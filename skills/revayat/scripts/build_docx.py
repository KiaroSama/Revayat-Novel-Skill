"""Stage 6 — build the Persian Word document from the translated Book IR.

Deliberately no Markdown and no Pandoc in this path. The IR already holds
everything a Word file needs — heading levels, emphasis spans, image bytes with
their physical size, footnote bodies — so routing it through Markdown would only
throw geometry away and then try to guess it back.

What the reader gets: real ``Heading N`` styles with bookmarks, a TOC field
whose entries are clickable, Word-native footnotes at the foot of the page,
pictures at their original size and aspect, right-to-left paragraphs with Latin
names left-to-right inside them, and selectable, editable Persian text.

The honest limit: Word reflows. A Persian paragraph is rarely the same length as
its English original, so page-for-page identity with the source PDF is not
achievable in an editable document, and this builder does not pretend to it.
Everything else on that list is exact.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt

import bookir as ir
import ooxml

#: EPUB and other reflowable sources give no physical size; assume screen DPI.
ASSUMED_DPI = 96.0

_LATIN_SEGMENT = re.compile(r"[A-Za-z](?:[A-Za-z0-9 .,'’&/@:_-]*[A-Za-z0-9])?")

_STYLE_BY_TYPE = {
    "paragraph": "Normal",
    "blockquote": "Quote",
    "caption": "Caption",
    "verse": "Normal",
}


def split_by_script(text: str) -> list[tuple[str, bool]]:
    """Split prose into ``(chunk, is_latin)`` segments.

    Persian runs are marked ``w:rtl``; Latin ones are not. Word then applies the
    Unicode bidi algorithm with the correct per-run character direction, which
    is why "الیزابت بنت (Elizabeth Bennet) رفت" comes out in the right order —
    and why nothing anywhere in this project reverses a string.

    Note that a run boundary is *not* a bidi boundary: brackets and other
    neutral characters around a Latin run resolve against the paragraph
    direction no matter which ``w:r`` holds them. Moving them between runs
    changes nothing on the page (measured in Word), so this deliberately does
    not try. Controlling that grouping would need Unicode isolates in the text
    itself, which is not worth putting invisible control characters into a
    manuscript people will edit.
    """
    segments: list[tuple[str, bool]] = []
    cursor = 0
    for match in _LATIN_SEGMENT.finditer(text):
        if match.start() > cursor:
            segments.append((text[cursor:match.start()], False))
        segments.append((match.group(0), True))
        cursor = match.end()
    if cursor < len(text):
        segments.append((text[cursor:], False))
    return [(chunk, latin) for chunk, latin in segments if chunk]


class Builder:
    def __init__(self, book: dict[str, Any], assets: Path, options: argparse.Namespace):
        self.book = book
        self.assets = assets
        self.options = options
        self.document = (
            Document(options.template) if options.template else Document()
        )
        self.footnotes = ooxml.Footnotes(self.document)
        self.bookmarks = ooxml.Bookmarks()
        self.notes_by_id = {n["id"]: n for n in book.get("footnotes", [])}
        self.used_notes: set[str] = set()
        self.warnings: list[str] = []
        self.toc_entries: list[tuple[str, str, int]] = []

    # -- document furniture ------------------------------------------------- #

    def setup(self) -> None:
        page = self.book.get("page", ir.default_page_setup())
        section = self.document.sections[0]
        section.page_width = Pt(page["width_pt"])
        section.page_height = Pt(page["height_pt"])
        section.top_margin = Pt(page["margin_top_pt"])
        section.bottom_margin = Pt(page["margin_bottom_pt"])
        section.left_margin = Pt(page["margin_inner_pt"])
        section.right_margin = Pt(page["margin_outer_pt"])
        ooxml.set_section_rtl(section, self.options.rtl)

        ooxml.set_document_defaults(
            self.document,
            persian_font=self.options.font,
            latin_font=self.options.latin_font,
            size_pt=self.options.size,
        )
        ooxml.ensure_footnote_styles(
            self.document, persian_font=self.options.font,
            size_pt=max(7.5, self.options.size - 2),
        )
        if self.options.rtl:
            for name in ("Normal", "Quote", "Caption", "List Bullet", "List Number",
                         *(f"Heading {n}" for n in range(1, 7))):
                ooxml.style_rtl(self.document, name, persian_font=self.options.font)
        ooxml.request_field_update(self.document)

    @property
    def text_width_pt(self) -> float:
        page = self.book.get("page", ir.default_page_setup())
        return max(
            72.0,
            page["width_pt"] - page["margin_inner_pt"] - page["margin_outer_pt"],
        )

    # -- writing ------------------------------------------------------------ #

    def paragraph(self, style: str = "Normal", *, align=None):
        paragraph = self.document.add_paragraph(style=style)
        ooxml.set_paragraph_rtl(paragraph, self.options.rtl)
        if align is not None:
            paragraph.alignment = align
        return paragraph

    def write_markup(self, paragraph, markup: str) -> None:
        """Emit one marked-up string as correctly-directioned Word runs."""
        for span in ir.parse_markup(markup):
            note_id = span.get("footnote")
            if note_id:
                self._write_footnote(paragraph, note_id)
                continue
            if span["verbatim"]:
                self._run(paragraph, span["text"], span, latin=True)
                continue
            for chunk, is_latin in split_by_script(span["text"]):
                self._run(paragraph, chunk, span, latin=is_latin)

    def _run(self, paragraph, text: str, span: dict[str, Any], *, latin: bool):
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bool(span.get("bold"))
        run.italic = bool(span.get("italic"))
        if latin:
            run.font.name = self.options.latin_font
        ooxml.set_run_direction(run, rtl=self.options.rtl and not latin)
        return run

    def _write_footnote(self, paragraph, note_id: str) -> None:
        note = self.notes_by_id.get(note_id)
        if note is None:
            self.warnings.append(f"footnote {note_id} referenced but not defined")
            return
        body = (note.get("target") or note.get("text") or "").strip()
        if not body:
            self.warnings.append(f"footnote {note_id} has no text")
            return
        spans: list[dict[str, Any]] = []
        for span in ir.parse_markup(body):
            if span.get("footnote"):
                continue  # nested footnotes are not a Word concept
            if span["verbatim"]:
                spans.append({**span, "verbatim": True})
                continue
            for chunk, is_latin in split_by_script(span["text"]):
                spans.append({**span, "text": chunk, "verbatim": is_latin})
        self.footnotes.add(paragraph, spans, persian_font=self.options.font,
                           rtl=self.options.rtl)
        self.used_notes.add(note_id)

    # -- blocks ------------------------------------------------------------- #

    def front_matter(self) -> None:
        meta = self.book.get("meta", {})
        title = (meta.get("title_target") or meta.get("title") or "").strip()
        author = (meta.get("author_target") or meta.get("author") or "").strip()
        if not title:
            return
        paragraph = self.paragraph("Title", align=WD_ALIGN_PARAGRAPH.CENTER)
        self.write_markup(paragraph, title)
        if author:
            byline = self.paragraph("Subtitle" if _has_style(self.document, "Subtitle")
                                    else "Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
            self.write_markup(byline, author)

    def table_of_contents(self) -> None:
        if not self.options.toc or not self.toc_entries:
            return
        heading = self.paragraph("Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)
        ooxml.page_break_before(heading)
        self.write_markup(heading, self.options.toc_title)
        holder = self.paragraph("Normal")
        ooxml.add_toc_field(holder, self.toc_entries, depth=self.options.toc_depth,
                            rtl=self.options.rtl)

    def collect_toc(self) -> None:
        """Anchor names are assigned before writing so the TOC can be first."""
        index = 0
        for block in self.book.get("blocks", []):
            if block["type"] != "heading":
                continue
            level = int(block.get("level", 1))
            if level > self.options.toc_depth:
                continue
            index += 1
            anchor = f"rv_{index:04d}"
            block["_anchor"] = anchor
            text = ir.plain_text(block.get("target") or block.get("text") or "")
            self.toc_entries.append((anchor, text.strip(), level))

    def body(self) -> None:
        first_heading = True
        for block in self.book.get("blocks", []):
            kind = block["type"]

            if kind == "heading":
                self._heading(block, first_heading)
                first_heading = False
            elif kind == "image":
                self._image(block)
            elif kind == "pagebreak":
                self._pagebreak(block)
            elif kind == "separator":
                paragraph = self.paragraph("Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
                self.write_markup(paragraph, "❖")
            elif kind in ir.TEXT_TYPES:
                self._text_block(block)

    def _heading(self, block: dict[str, Any], first: bool) -> None:
        level = min(6, max(1, int(block.get("level", 1))))
        paragraph = self.paragraph(
            f"Heading {level}",
            align=WD_ALIGN_PARAGRAPH.CENTER if level == 1 else None,
        )
        if level == 1 and self.options.page_breaks != "none" and not first:
            ooxml.page_break_before(paragraph)
        ooxml.keep_with_next(paragraph)
        anchor = block.get("_anchor")
        text = self._text_of(block)
        self.write_markup(paragraph, text)
        self._apply_source_size(paragraph, block)
        if anchor:
            self.bookmarks.wrap(paragraph, anchor, ir.plain_text(text))

    def _apply_source_size(self, paragraph, block: dict[str, Any]) -> None:
        """Reproduce the heading's size from the source book, when asked.

        Extraction records the real point size of every heading it finds. The
        default still uses the Word ``Heading N`` styles — they give a coherent
        document even when the source's own sizes are erratic — but
        ``--heading-size source`` reinstates the book's exact metrics, which is
        what "keep the book's styling" actually means.
        """
        if self.options.heading_size != "source":
            return
        size = block.get("font_size_pt")
        if not size:
            return
        for run in paragraph.runs:
            run.font.size = Pt(float(size))

    def _text_block(self, block: dict[str, Any]) -> None:
        text = self._text_of(block)
        if not text.strip():
            return
        kind = block["type"]
        if kind == "listitem":
            style = "List Number" if block.get("ordered") else "List Bullet"
            style = style if _has_style(self.document, style) else "Normal"
        else:
            style = _STYLE_BY_TYPE.get(kind, "Normal")
            style = style if _has_style(self.document, style) else "Normal"

        align = None
        if kind == "paragraph" and self.options.justify:
            align = WD_ALIGN_PARAGRAPH.JUSTIFY
        elif kind in ("caption", "verse"):
            align = WD_ALIGN_PARAGRAPH.CENTER

        paragraph = self.paragraph(style, align=align)
        self.write_markup(paragraph, text)

    def _text_of(self, block: dict[str, Any]) -> str:
        target = (block.get("target") or "").strip()
        if target:
            return target
        source = (block.get("text") or "").strip()
        if source:
            self.warnings.append(f"block {block['id']} is untranslated")
        return source

    def _image(self, block: dict[str, Any]) -> None:
        path = self.assets / block["asset"]
        if not path.exists():
            self.warnings.append(f"missing asset {block['asset']}")
            return

        width_pt, height_pt = self._image_size(block)
        paragraph = self.paragraph("Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
        run = paragraph.add_run()
        try:
            if height_pt:
                run.add_picture(str(path), width=Pt(width_pt), height=Pt(height_pt))
            else:
                run.add_picture(str(path), width=Pt(width_pt))
        except Exception as error:  # unsupported or corrupt image stream
            self.warnings.append(f"could not place {block['asset']}: {error}")
            return

        alt = (block.get("target_alt") or block.get("alt") or "").strip()
        if alt and self.options.captions:
            caption = self.paragraph(
                "Caption" if _has_style(self.document, "Caption") else "Normal",
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            self.write_markup(caption, alt)

    def _image_size(self, block: dict[str, Any]) -> tuple[float, float | None]:
        """Physical size in points, preserving aspect and fitting the text block."""
        available = self.text_width_pt
        width = block.get("width_pt")
        height = block.get("height_pt")

        if not width:
            pixel_width = block.get("pixel_width")
            pixel_height = block.get("pixel_height")
            if pixel_width:
                width = pixel_width * 72.0 / ASSUMED_DPI
                height = (pixel_height * 72.0 / ASSUMED_DPI) if pixel_height else None
            else:
                return available * 0.6, None

        if width > available:
            if height:
                height = height * (available / width)   # aspect ratio is preserved
            width = available
        return width, height

    def _pagebreak(self, block: dict[str, Any]) -> None:
        if self.options.page_breaks != "source" or block.get("soft"):
            return
        paragraph = self.paragraph("Normal")
        paragraph.add_run().add_break(WD_BREAK.PAGE)

    # -- run ---------------------------------------------------------------- #

    def build(self, destination: Path) -> dict[str, Any]:
        self.setup()
        self.collect_toc()
        self.front_matter()
        self.table_of_contents()
        self.body()
        self.footnotes.finalise()

        unused = [n for n in self.notes_by_id if n not in self.used_notes]
        if unused:
            self.warnings.append(
                f"{len(unused)} footnote(s) never referenced: {', '.join(unused[:5])}"
            )

        destination.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(destination))
        return {
            "output": str(destination),
            "headings": len(self.toc_entries),
            "bookmarks": len(self.bookmarks.names),
            "footnotes": len(self.footnotes),
            "images": sum(1 for b in self.book["blocks"] if b["type"] == "image"),
            "warnings": self.warnings[:40],
            "warning_count": len(self.warnings),
        }


def _has_style(document, name: str) -> bool:
    try:
        document.styles[name]
        return True
    except KeyError:
        return False


def add_arguments(parser: argparse.ArgumentParser) -> None:
    parser.add_argument("--book", required=True)
    parser.add_argument("--assets", default=None,
                        help="asset directory (default: <book dir>/assets)")
    parser.add_argument("--out", required=True, help="output .docx path")
    parser.add_argument("--template", default=None,
                        help="reference .docx supplying styles and page setup")
    parser.add_argument("--font", default="Vazirmatn",
                        help="Persian (complex-script) font; e.g. 'B Nazanin', 'Tahoma'")
    parser.add_argument("--latin-font", default="Times New Roman")
    parser.add_argument("--size", type=float, default=11.5, help="body size in pt")
    parser.add_argument("--toc", action=argparse.BooleanOptionalAction, default=True)
    parser.add_argument("--toc-title", default="فهرست مطالب")
    parser.add_argument("--toc-depth", type=int, default=2)
    parser.add_argument("--justify", action=argparse.BooleanOptionalAction, default=True)
    parser.add_argument("--captions", action=argparse.BooleanOptionalAction, default=True)
    parser.add_argument("--rtl", action=argparse.BooleanOptionalAction, default=True)
    parser.add_argument("--page-breaks", choices=["chapter", "source", "none"],
                        default="chapter")
    parser.add_argument("--heading-size", choices=["style", "source"], default="style",
                        help="'style' uses Word's Heading N sizes; 'source' reproduces "
                             "the point size measured in the original book")


def main(argv: list[str] | None = None) -> int:
    ir.use_utf8_stdio()
    parser = argparse.ArgumentParser(prog="revayat build", description=__doc__)
    add_arguments(parser)
    args = parser.parse_args(argv)

    book_path = Path(args.book)
    book = ir.load_book(book_path)
    assets = Path(args.assets) if args.assets else book_path.parent / "assets"

    report = Builder(book, assets, args).build(Path(args.out))
    print(json.dumps(report, ensure_ascii=False, indent=1))
    return 0


if __name__ == "__main__":
    sys.exit(main())
