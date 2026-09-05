"""The Persian publication profile — page and paragraph layout, set once.

Everything here is applied to Word **styles**, not to individual paragraphs.
That is the difference between a document a designer can restyle in one click
and one where every paragraph carries its own hard-coded spacing: change
``Normal`` and the whole book follows. Direct formatting is reserved for the few
things that genuinely vary per paragraph — alignment on a centred heading, a
picture's size.

The defaults are trade-book conventions rather than Word's: justified body with
a first-line indent and no blank line between paragraphs, which is how Persian
prose is actually set. Word's default — ragged, extra space between paragraphs,
no indent — reads as a report, not a book.

Nothing here can be inferred from the source PDF, so none of it pretends to be
"the original layout". It is a deliberate house style, and every part of it is
adjustable from the command line.
"""

from __future__ import annotations

from typing import Any

from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn

W = nsmap["w"]
_WNS = f'xmlns:w="{W}"'

#: Word measures spacing in twentieths of a point.
TWIPS_PER_POINT = 20


def twips(points: float) -> int:
    return int(round(points * TWIPS_PER_POINT))


class Profile:
    """Layout values for one book. All lengths are points unless stated."""

    def __init__(
        self,
        *,
        line_spacing: float = 1.5,
        first_line_indent_pt: float = 18.0,
        space_before_pt: float = 0.0,
        space_after_pt: float = 0.0,
        heading_space_before_pt: float = 18.0,
        heading_space_after_pt: float = 10.0,
        caption_space_after_pt: float = 12.0,
        quote_indent_pt: float = 24.0,
        widow_control: bool = True,
        mirror_margins: bool = False,
        gutter_pt: float = 0.0,
        header_distance_pt: float = 36.0,
        footer_distance_pt: float = 36.0,
        page_numbers: bool = True,
    ) -> None:
        self.line_spacing = line_spacing
        self.first_line_indent_pt = first_line_indent_pt
        self.space_before_pt = space_before_pt
        self.space_after_pt = space_after_pt
        self.heading_space_before_pt = heading_space_before_pt
        self.heading_space_after_pt = heading_space_after_pt
        self.caption_space_after_pt = caption_space_after_pt
        self.quote_indent_pt = quote_indent_pt
        self.widow_control = widow_control
        self.mirror_margins = mirror_margins
        self.gutter_pt = gutter_pt
        self.header_distance_pt = header_distance_pt
        self.footer_distance_pt = footer_distance_pt
        self.page_numbers = page_numbers

    @classmethod
    def from_options(cls, options: Any) -> "Profile":
        return cls(
            line_spacing=getattr(options, "line_spacing", 1.5),
            first_line_indent_pt=getattr(options, "first_line_indent", 18.0),
            widow_control=getattr(options, "widow_control", True),
            mirror_margins=getattr(options, "mirror_margins", False),
            gutter_pt=getattr(options, "gutter", 0.0),
            page_numbers=getattr(options, "page_numbers", True),
        )


# --------------------------------------------------------------------------- #
# Style-level formatting
# --------------------------------------------------------------------------- #

def _paragraph_properties(style_element):
    properties = style_element.find(qn("w:pPr"))
    if properties is None:
        properties = parse_xml(f"<w:pPr {_WNS}/>")
        style_element.append(properties)
    return properties


def _set(properties, tag: str, **attrs: Any) -> None:
    """Replace a single paragraph-property element, keeping the rest."""
    for existing in properties.findall(qn(f"w:{tag}")):
        properties.remove(existing)
    rendered = "".join(f' w:{name}="{value}"' for name, value in attrs.items())
    properties.append(parse_xml(f"<w:{tag} {_WNS}{rendered}/>"))


def _style(document, name: str):
    try:
        return document.styles[name].element
    except KeyError:
        return None


def apply_paragraph_styles(document, profile: Profile, *, rtl: bool = True) -> list[str]:
    """Write the profile into the document's styles. Returns the ones touched."""
    touched: list[str] = []

    body = _style(document, "Normal")
    if body is not None:
        properties = _paragraph_properties(body)
        _set(properties, "spacing",
             before=twips(profile.space_before_pt),
             after=twips(profile.space_after_pt),
             line=twips(profile.line_spacing * 12), lineRule="auto")
        # In an RTL paragraph the first-line indent is still w:firstLine; Word
        # applies it to the reading-order start, which is the right-hand edge.
        _set(properties, "ind", firstLine=twips(profile.first_line_indent_pt))
        if profile.widow_control:
            _set(properties, "widowControl", val="1")
        touched.append("Normal")

    for level in range(1, 7):
        heading = _style(document, f"Heading {level}")
        if heading is None:
            continue
        properties = _paragraph_properties(heading)
        _set(properties, "spacing",
             before=twips(profile.heading_space_before_pt),
             after=twips(profile.heading_space_after_pt),
             line=twips(profile.line_spacing * 12), lineRule="auto")
        # A heading never carries a first-line indent, and must never be the
        # last line on a page.
        _set(properties, "ind", firstLine=0)
        _set(properties, "keepNext", val="1")
        _set(properties, "keepLines", val="1")
        touched.append(f"Heading {level}")

    quote = _style(document, "Quote")
    if quote is not None:
        properties = _paragraph_properties(quote)
        _set(properties, "spacing", before=twips(6), after=twips(6),
             line=twips(profile.line_spacing * 12), lineRule="auto")
        # Logical indent: "start" is the right edge in an RTL paragraph.
        _set(properties, "ind", start=twips(profile.quote_indent_pt), firstLine=0)
        _set(properties, "keepLines", val="1")
        touched.append("Quote")

    caption = _style(document, "Caption")
    if caption is not None:
        properties = _paragraph_properties(caption)
        _set(properties, "spacing", before=twips(4),
             after=twips(profile.caption_space_after_pt), line=twips(12),
             lineRule="auto")
        _set(properties, "ind", firstLine=0)
        # A caption belongs to the picture above it.
        _set(properties, "keepLines", val="1")
        touched.append("Caption")

    return touched


# --------------------------------------------------------------------------- #
# Section-level formatting
# --------------------------------------------------------------------------- #

def apply_section(section, profile: Profile) -> None:
    """Mirrored margins, gutter and header/footer distances."""
    properties = section._sectPr

    if profile.mirror_margins:
        # Mirrored margins make the *inside* margin the binding edge on both
        # sides, which is what a printed book needs and what "inner/outer"
        # in the page setup actually means.
        for existing in properties.findall(qn("w:mirrorMargins")):
            properties.remove(existing)
        properties.insert(0, parse_xml(f'<w:mirrorMargins {_WNS}/>'))

    page_margin = properties.find(qn("w:pgMar"))
    if page_margin is not None:
        page_margin.set(qn("w:gutter"), str(twips(profile.gutter_pt)))
        page_margin.set(qn("w:header"), str(twips(profile.header_distance_pt)))
        page_margin.set(qn("w:footer"), str(twips(profile.footer_distance_pt)))


def add_page_numbers(document, section, *, rtl: bool = True) -> bool:
    """A centred page number in the footer, as a live PAGE field.

    A field rather than literal text, so the numbers stay correct when the
    Persian reflows to a different number of pages than the source.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    for run in list(paragraph.runs):
        run._r.getparent().remove(run._r)

    properties = paragraph._p.get_or_add_pPr()
    if rtl:
        for existing in properties.findall(qn("w:bidi")):
            properties.remove(existing)
        properties.append(parse_xml(f"<w:bidi {_WNS}/>"))
    _set(properties, "jc", val="center")

    digits = '<w:rPr><w:rtl/><w:lang w:bidi="fa-IR"/></w:rPr>' if rtl else ""
    for fragment in (
        f'<w:r {_WNS}>{digits}<w:fldChar w:fldCharType="begin"/></w:r>',
        f'<w:r {_WNS}>{digits}<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>',
        f'<w:r {_WNS}>{digits}<w:fldChar w:fldCharType="separate"/></w:r>',
        f'<w:r {_WNS}>{digits}<w:t>1</w:t></w:r>',
        f'<w:r {_WNS}>{digits}<w:fldChar w:fldCharType="end"/></w:r>',
    ):
        paragraph._p.append(parse_xml(fragment))
    return True


def apply(document, section, profile: Profile, *, rtl: bool = True) -> dict[str, Any]:
    """Apply the whole profile. Returns what was done, for the build report."""
    styles = apply_paragraph_styles(document, profile, rtl=rtl)
    apply_section(section, profile)
    numbered = add_page_numbers(document, section, rtl=rtl) if profile.page_numbers else False
    return {
        "styles": styles,
        "mirror_margins": profile.mirror_margins,
        "gutter_pt": profile.gutter_pt,
        "line_spacing": profile.line_spacing,
        "first_line_indent_pt": profile.first_line_indent_pt,
        "page_numbers": numbered,
    }
