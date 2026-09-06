"""DOCX → Book IR via python-docx, plus raw OOXML for what it cannot reach.

python-docx exposes paragraphs, runs and styles, but has no API for footnotes
(verified against python-docx 1.2.0), so ``word/footnotes.xml`` is read
straight off the package. Inline pictures are pulled from their relationship
so the original bytes and the author's ``wp:extent`` both survive.

Section breaks travel as ``book["sections"]``, each entry naming the block it
opens at. ``book["page"]`` still reports the first section's geometry, exactly
as it did when that was the only geometry the reader kept.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterator

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.hyperlink import Hyperlink
from docx.text.paragraph import Paragraph
from docx.text.run import Run

import bookir as ir

EMU_PER_PT = ir.EMU_PER_PT

_HEADING_STYLE = re.compile(r"^\s*heading\s*(\d)\s*$", re.I)
_LIST_STYLE = re.compile(r"list\s*(bullet|number|paragraph)", re.I)


def column_span(tc) -> int:
    """How many grid columns this ``w:tc`` covers, from ``w:gridSpan``."""
    properties = tc.find(qn("w:tcPr"))
    if properties is None:
        return 1
    grid = properties.find(qn("w:gridSpan"))
    try:
        return max(1, int(grid.get(qn("w:val"))))
    except (AttributeError, TypeError, ValueError):
        return 1


def _vertical_merge(tc) -> str | None:
    """``"restart"``, ``"continue"``, or ``None`` when the cell is not merged."""
    properties = tc.find(qn("w:tcPr"))
    if properties is None:
        return None
    merge = properties.find(qn("w:vMerge"))
    if merge is None:
        return None
    # A bare <w:vMerge/> means continue; only "restart" opens a new span.
    return "restart" if merge.get(qn("w:val")) == "restart" else "continue"


def table_cells(table) -> list[dict[str, Any]]:
    """Every distinct cell of a table, once, with its position and span.

    Walks ``w:tr``/``w:tc`` rather than ``row.cells``. Two reasons, and the
    second one is the reason this is not shorter:

    ``row.cells`` *expands* merges — a cell spanning two columns comes back
    twice, and a vertically merged one comes back on every row it covers. Since
    every cell here becomes its own worksheet unit, that made a merged cell's
    sentence get translated twice and printed twice.

    And de-duplicating that expansion by object identity does not work.
    `cell._tc` hands back a fresh lxml proxy on each access, and CPython reuses
    the `id()` of a freed one — measured on a 3x3 grid, three different cells
    reported the same id and a fourth reported one that had never been seen.
    Walking the XML gives each cell exactly once by construction, so there is
    nothing to de-duplicate.

    Returns ``{"tc", "row", "cell", "row_span", "col_span"}`` with 1-based
    ``row``/``cell`` counted in grid columns.
    """
    from docx.table import _Cell

    found: list[dict[str, Any]] = []
    # Grid column -> the record of the cell currently open there, so a
    # `continue` row extends the span of the cell that started it.
    open_at: dict[int, dict[str, Any]] = {}

    for row_number, tr in enumerate(table._tbl.findall(qn("w:tr")), start=1):
        column = 0
        for tc in tr.findall(qn("w:tc")):
            width = column_span(tc)
            if _vertical_merge(tc) == "continue" and column in open_at:
                open_at[column]["row_span"] += 1
            else:
                record = {"tc": _Cell(tc, table), "row": row_number,
                          "cell": column + 1, "row_span": 1, "col_span": width}
                found.append(record)
                open_at[column] = record
            column += width
    return found


def iter_runs(paragraph: Paragraph) -> Iterator[Run]:
    """Every run in the paragraph, including the ones inside a hyperlink.

    ``paragraph.runs`` omits hyperlink content entirely. A sentence with a
    linked phrase in the middle of it therefore came through with the phrase
    missing — not flagged, not empty, just quietly shorter than the source.
    """
    for item in paragraph.iter_inner_content():
        if isinstance(item, Run):
            yield item
        elif isinstance(item, Hyperlink):
            yield from item.runs


def hyperlinks(paragraph: Paragraph) -> list[dict[str, str]]:
    """The links in this paragraph, as ``{"text", "href"}``.

    ``iter_runs`` keeps a link's *words* - that was the loss worth fixing first,
    because a missing phrase reads as a sentence the author wrote that way. The
    target is the other half: it survives here, on the block, rather than in the
    markup the translator answers. A URL is not text to translate, and a paired
    inline marker for something a printed Persian book cannot click would be a
    contract every worksheet has to honour for no reader's benefit.
    """
    found = []
    for item in paragraph.iter_inner_content():
        if not isinstance(item, Hyperlink):
            continue
        # `address` is the external target; `fragment` the in-document anchor.
        # A link may have either or both, and one without the other is normal.
        target = item.address or ""
        if item.fragment:
            target = f"{target}#{item.fragment}" if target else f"#{item.fragment}"
        if item.text and target:
            found.append({"text": item.text, "href": target})
    return found


def has_page_break(run: Run) -> bool:
    return any(node.get(qn("w:type")) == "page"
               for node in run._r.findall(qn("w:br")))


def _pt(length) -> float | None:
    """Points, rounded like the rest of the page setup; ``None`` when unset."""
    return None if length is None else round(length.pt, 2)


def section_geometry(section) -> dict[str, Any]:
    """One ``w:sectPr`` as page setup, in the field names ``book["page"]`` uses.

    The names match deliberately: the first section *is* ``book["page"]``, so
    everything that already reads that key keeps working while the rest of the
    sections travel beside it in the same shape.

    ``start_type`` is kept as the raw XML token rather than an enum member so
    the value survives a JSON round trip and goes back into the built document
    unchanged.
    """
    start = section._sectPr.find(qn("w:type"))
    token = start.get(qn("w:val")) if start is not None else None
    return {
        "start_type": token or "nextPage",   # Word's default when w:type is absent
        "orientation": ("landscape" if section.orientation == WD_ORIENT.LANDSCAPE
                        else "portrait"),
        "width_pt": _pt(section.page_width),
        "height_pt": _pt(section.page_height),
        "margin_top_pt": _pt(section.top_margin),
        "margin_bottom_pt": _pt(section.bottom_margin),
        "margin_inner_pt": _pt(section.left_margin),
        "margin_outer_pt": _pt(section.right_margin),
        "gutter_pt": _pt(section.gutter),
        "header_distance_pt": _pt(section.header_distance),
        "footer_distance_pt": _pt(section.footer_distance),
    }


def ends_section(paragraph: Paragraph) -> bool:
    """True when this paragraph carries the ``w:sectPr`` that closes a section.

    A section's properties sit at its *end*, not its start: on the last
    paragraph of the section for all but the last one, and on ``w:body`` for
    that. So the block after this paragraph is the first of the next section.
    """
    properties = paragraph._p.find(qn("w:pPr"))
    return properties is not None and properties.find(qn("w:sectPr")) is not None


def _column_count(section) -> int:
    columns = section._sectPr.find(qn("w:cols"))
    try:
        return max(1, int(columns.get(qn("w:num"))))
    except (AttributeError, TypeError, ValueError):
        return 1


def _ilvl(element) -> int | None:
    """The 0-based ``w:ilvl`` under this element, as a 1-based depth."""
    if element is None:
        return None
    numbering = element.find(f".//{qn('w:numPr')}")
    if numbering is None:
        return None
    level = numbering.find(qn("w:ilvl"))
    try:
        return int(level.get(qn("w:val"))) + 1
    except (AttributeError, TypeError, ValueError):
        return 1


def list_level(paragraph: Paragraph) -> int | None:
    """The nesting depth of a list item, 1-based, or ``None`` if not a list.

    Word records the depth in one of two places and it is genuinely either:
    a paragraph numbered directly carries ``w:numPr/w:ilvl`` itself, while one
    numbered through a style — which is what "List Bullet 2" is — carries
    nothing, and the depth lives in the style definition.

    When the style says a paragraph is a list but neither place gives a level,
    the trailing digit of the style name is used. That is a reading of Word's
    own naming convention rather than of the file's data, so it is last, and it
    is only ever reached for a paragraph already known to be a list item.
    """
    depth = _ilvl(paragraph._p.find(qn("w:pPr")))
    if depth is not None:
        return depth

    style = getattr(paragraph, "style", None)
    element = getattr(style, "element", None)
    depth = _ilvl(element)
    if depth is not None and depth > 1:
        return depth

    name = (getattr(style, "name", "") or "").strip()
    if not _LIST_STYLE.search(name):
        return None
    trailing = re.search(r"(\d+)\s*$", name)
    return int(trailing.group(1)) if trailing else 1


def _heading_level(style_name: str) -> int | None:
    name = (style_name or "").strip()
    if name.lower() in {"title", "subtitle"}:
        return 1 if name.lower() == "title" else 2
    match = _HEADING_STYLE.match(name)
    if match:
        return min(6, max(1, int(match.group(1))))
    return None


def _run_style(run) -> tuple[bool, bool]:
    """Effective bold/italic, falling back to the run's style definition."""
    bold, italic = run.bold, run.italic
    if bold is None or italic is None:
        style = getattr(run, "style", None)
        font = getattr(style, "font", None) if style is not None else None
        if font is not None:
            bold = font.bold if bold is None else bold
            italic = font.italic if italic is None else italic
    return bool(bold), bool(italic)


#: The two ways Word stores a note. They are the same shape and the same loss
#: if missed - a book that puts its notes at the back rather than the foot of
#: the page came through with every one of them gone, and nothing said so.
#: Both become footnotes in the Persian edition, because that is where a
#: Persian reader looks, and because it is what the builder writes.
NOTE_PARTS = (("footnote", RT.FOOTNOTES, "w:footnote", "w:footnoteReference"),
              ("endnote", RT.ENDNOTES, "w:endnote", "w:endnoteReference"))


def _read_notes(document) -> dict[str, str]:
    """``"<kind>:<id>" -> plain text``, from both note parts.

    Keyed by kind as well as id because the two id spaces are separate: an
    endnote 1 and a footnote 1 are different notes, and merging them on the
    number alone silently drops one of the two.

    Word reserves ids -1 and 0 for the separator marks; they carry no content.
    """
    from lxml import etree

    notes: dict[str, str] = {}
    for kind, relationship, tag, _ in NOTE_PARTS:
        try:
            part = document.part.part_related_by(relationship)
        except KeyError:
            continue
        try:
            root = etree.fromstring(part.blob)
        except Exception:
            continue
        for node in root.findall(qn(tag)):
            note_id = node.get(qn("w:id"))
            if note_id is None or int(note_id) <= 0:
                continue
            pieces = [t.text or "" for t in node.iter(qn("w:t"))]
            text = re.sub(r"\s+", " ", "".join(pieces)).strip()
            if text:
                notes[f"{kind}:{note_id}"] = text
    return notes


def _picture(run, document) -> dict[str, Any] | None:
    """Image bytes and the author's rendered size, from an inline drawing."""
    blips = run._r.findall(f".//{qn('a:blip')}")
    if not blips:
        return None
    rel_id = blips[0].get(qn("r:embed"))
    if not rel_id:
        return None
    try:
        image_part = document.part.related_parts[rel_id]
    except KeyError:
        return None

    width_pt = height_pt = None
    extent = run._r.find(f".//{qn('wp:extent')}")
    if extent is not None:
        try:
            width_pt = round(int(extent.get("cx")) / EMU_PER_PT, 2)
            height_pt = round(int(extent.get("cy")) / EMU_PER_PT, 2)
        except (TypeError, ValueError):
            pass

    blob = image_part.blob
    return {
        "blob": blob,
        "name": Path(str(image_part.partname)).name,
        "sha256": ir.sha256_bytes(blob),
        "width_pt": width_pt,
        "height_pt": height_pt,
    }


def read_docx(
    path: str,
    asset_dir: Path,
    *,
    lang_source: str = "en",
    lang_target: str = "fa-IR",
) -> dict[str, Any]:
    asset_dir.mkdir(parents=True, exist_ok=True)
    document = Document(path)
    notes = _read_notes(document)
    core = document.core_properties

    book = ir.new_book(
        source_path=str(path),
        source_format="docx",
        source_sha256=ir.sha256_file(path),
        pages=0,
        title=(core.title or Path(path).stem).strip(),
        author=(core.author or "").strip(),
        lang_source=lang_source,
        lang_target=lang_target,
    )

    section = document.sections[0] if document.sections else None
    if section is not None and section.page_width and section.page_height:
        book["page"].update({
            "width_pt": round(section.page_width.pt, 2),
            "height_pt": round(section.page_height.pt, 2),
            "margin_top_pt": round(section.top_margin.pt, 2),
            "margin_bottom_pt": round(section.bottom_margin.pt, 2),
            "margin_inner_pt": round(section.left_margin.pt, 2),
            "margin_outer_pt": round(section.right_margin.pt, 2),
        })

    blocks: list[dict[str, Any]] = []
    footnotes: list[dict[str, Any]] = []
    seen_assets: dict[str, str] = {}
    used_notes: set[str] = set()
    counter = 0

    # Only the bookmarks an in-document link actually points at are carried.
    # A Word file is full of `_GoBack` and `_Toc…` names nothing refers to, and
    # every one of them would become a bookmark in the Persian edition for no
    # reader's benefit — while an anchor whose destination was dropped is a
    # dead link the package gate rightly refuses.
    anchored = {node.get(qn("w:anchor"))
                for node in document.element.body.iter(qn("w:hyperlink"))
                if node.get(qn("w:anchor"))}
    placed_anchors: set[str] = set()

    def add(block_type: str, **fields: Any) -> dict[str, Any]:
        nonlocal counter
        counter += 1
        block = ir.make_block(block_type, counter, **fields)
        blocks.append(block)
        return block

    warnings: list[dict[str, Any]] = []

    def keep_anchors(paragraph, first: int) -> None:
        """File the paragraph's linked-to bookmarks on the block it became.

        On the *first* block only. A paragraph split by a picture or a page
        break becomes several blocks, and repeating the name on each one would
        open the same bookmark two or three times — which Word resolves by
        sending every link to the first, silently.
        """
        names = [name for node in paragraph._p.iter(qn("w:bookmarkStart"))
                 if (name := node.get(qn("w:name"))) in anchored
                 and name not in placed_anchors]
        if not names:
            return
        block = next((b for b in blocks[first:] if b["type"] in ir.TEXT_TYPES), None)
        if block is None:
            return
        block["bookmarks"] = names
        placed_anchors.update(names)

    def read_paragraph(paragraph, **extra: Any) -> None:
        first_block = len(blocks)
        spans: list[tuple[str, bool, bool]] = []
        pending_notes: list[str] = []
        # A link belongs to the prose of this paragraph, never to a picture
        # that happens to sit inside it, so the two branches carry different
        # context.
        links = hyperlinks(paragraph)
        prose = {**extra, "links": links} if links else extra

        for run in iter_runs(paragraph):
            picture = _picture(run, document)
            if picture is not None:
                _flush(spans, pending_notes, paragraph, add, footnotes,
                       used_notes, notes, prose)
                spans, pending_notes = [], []
                digest = picture["sha256"]
                if digest in seen_assets:
                    asset_name = seen_assets[digest]
                else:
                    asset_name = f"d{len(seen_assets) + 1:04d}-{picture['name']}"
                    (asset_dir / asset_name).write_bytes(picture["blob"])
                    seen_assets[digest] = asset_name
                add("image", page=0, asset=asset_name, sha256=digest, bbox=None,
                    width_pt=picture["width_pt"], height_pt=picture["height_pt"],
                    pixel_width=None, pixel_height=None, alt="", target_alt=None,
                    **extra)
                continue

            for kind, _, _, reference in NOTE_PARTS:
                for ref in run._r.findall(qn(reference)):
                    key = f"{kind}:{ref.get(qn('w:id'))}"
                    if key in notes:
                        pending_notes.append(key)

            text = run.text
            if text:
                bold, italic = _run_style(run)
                spans.append((text, bold, italic))

            if has_page_break(run):
                _flush(spans, pending_notes, paragraph, add, footnotes,
                       used_notes, notes, prose)
                spans, pending_notes = [], []
                add("pagebreak", page=0, soft=False)

        _flush(spans, pending_notes, paragraph, add, footnotes, used_notes,
               notes, prose)
        keep_anchors(paragraph, first_block)

    def read_table(table, table_id: str, depth: int = 0) -> int:
        """One table's cells, each read exactly once. Returns the row count.

        Two traps, both of which used to be live.

        ``row.cells`` *expands* merges: a cell spanning two columns comes back
        twice, and a vertically merged one comes back on every row it covers —
        the same `w:tc` object each time. Since every cell becomes its own
        worksheet unit, that meant a merged cell's sentence was translated
        twice and printed twice. Cells are therefore keyed by the identity of
        the underlying element, and the span is recorded instead.

        And a table nested in a cell is not in ``document.iter_inner_content``
        at all, so its text vanished exactly as every table's did before.
        """
        rows = 0
        merges = 0
        for record in table_cells(table):
            rows = max(rows, record["row"] + record["row_span"] - 1)
            span = {name: record[name] for name in ("row_span", "col_span")
                    if record[name] > 1}
            merges += bool(span)
            cell = record["tc"]
            for paragraph in cell.paragraphs:
                read_paragraph(paragraph, table=table_id, row=record["row"],
                               cell=record["cell"], **span)
            for inner_number, inner in enumerate(cell.tables, start=1):
                read_table(inner, f"{table_id}-{record['row']}"
                                  f"{record['cell']}n{inner_number}", depth + 1)
        if merges:
            warnings.append({
                "kind": "table-merged-cells", "table": table_id,
                "detail": f"{merges} merged cell position(s); the text is read "
                          f"once and its span recorded, and the builder "
                          f"reproduces the merge",
            })
        return rows

    # The body in document order. `document.paragraphs` walks only top-level
    # paragraphs, so every table in the book — every cell of it — was dropped
    # without a word: not flagged, not empty, simply absent from the IR and
    # therefore from the translation and from the finished file.
    #
    # `section_starts` is where each section begins, as a count of blocks read
    # so far: section 0 opens the book, and every later one opens right after
    # the paragraph carrying the previous section's `w:sectPr`.
    section_starts = [0]
    for index, item in enumerate(document.iter_inner_content()):
        if isinstance(item, Paragraph):
            read_paragraph(item)
            if ends_section(item):
                section_starts.append(len(blocks))
        elif isinstance(item, Table):
            read_table(item, f"t{index:04d}")

    book["blocks"] = blocks
    book["footnotes"] = footnotes

    def start_block(number: int) -> str | None:
        """The id of the first block of section ``number``.

        ``None`` when the section holds no blocks at all — a document that ends
        with a section break has a final section made of nothing, and there is
        nothing in the built file for the break to sit before.
        """
        if number >= len(section_starts):
            return None
        at = section_starts[number]
        return blocks[at]["id"] if at < len(blocks) else None

    book["sections"] = [
        {"index": number, "start_block": start_block(number),
         **section_geometry(section)}
        for number, section in enumerate(document.sections)
    ]
    linked = sum(len(block.get("links") or []) for block in blocks)
    if linked:
        warnings.append({
            "kind": "hyperlinks-kept-as-metadata", "count": linked,
            "detail": "link text is in the prose and each target is on its "
                      "block as `links`; the builder puts a live link back "
                      "only where the translation kept the display phrase "
                      "word for word, and names every one it could not",
        })

    running = sum(1 for relationship in document.part.rels.values()
                  if relationship.reltype in (RT.HEADER, RT.FOOTER))
    if running:
        warnings.append({
            "kind": "running-heads-dropped", "count": running,
            "detail": "the source's headers and footers are not read; the "
                      "built document generates its own from --page-numbers",
        })

    # `sections-collapsed` used to be reported here. It said the truth at the
    # time — only the first section's geometry was kept — and became a lie the
    # moment `book["sections"]` started carrying every one of them. What is
    # still genuinely dropped is named on its own below.
    for record, section in zip(book["sections"], document.sections):
        columns = _column_count(section)
        if columns > 1:
            warnings.append({
                "kind": "section-columns-dropped", "section": record["index"],
                "count": columns,
                "detail": f"section {record['index']} is set in {columns} "
                          f"columns; page size, orientation and margins are "
                          f"carried but the Persian edition is set in one",
            })

    if warnings:
        book["source"]["docx_warnings"] = warnings
    return book


def _flush(spans, pending_notes, paragraph, add, footnotes, used_notes, notes,
           extra: dict[str, Any] | None = None) -> None:
    """Emit the accumulated runs of one paragraph as a block."""
    if not spans:
        return
    text = ir.render_markup(spans).strip()
    if not text:
        return
    extra = dict(extra or {})

    for note_id in pending_notes:
        if note_id in used_notes:
            continue
        used_notes.add(note_id)
        note = ir.make_footnote(len(footnotes) + 1, anchor_block="",
                                text=notes[note_id], origin="source")
        footnotes.append(note)
        text = f"{text}[[fn:{note['id']}]]"

    style_name = getattr(paragraph.style, "name", "") or ""
    # The style records the paragraph's role; the numbering properties record
    # how deep it sits. A nested list needs both.
    depth = list_level(paragraph)
    level = _heading_level(style_name)
    if level is not None:
        block = add("heading", page=0, level=level, text=text, **extra)
    elif style_name.lower().startswith("quote") or style_name.lower() == "intense quote":
        block = add("blockquote", page=0, text=text, **extra)
    elif depth is not None or _LIST_STYLE.search(style_name):
        block = add("listitem", page=0, level=depth or 1,
                    ordered="number" in style_name.lower(), text=text, **extra)
    elif style_name.lower() == "caption":
        block = add("caption", page=0, text=text, **extra)
    else:
        block = add("paragraph", page=0, text=text, **extra)

    for note in footnotes:
        if not note["anchor_block"] and note["id"] in ir.footnote_refs(text):
            note["anchor_block"] = block["id"]
