"""Running heads and feet: carried across, and set in Persian.

They used to be counted and thrown away, and the `running-heads-dropped`
warning recorded why: an English title marching across a Persian page is worse
than no running head at all. That reasoning is intact. It is answered here by
translating the author's own head instead of discarding it, so these tests
assert against the *source document* — this Word file printed this line above
its prose, therefore the Persian one prints that line, in Persian — and against
the thing the old decision was really protecting, which is that no English
reaches a Persian page however the translation goes wrong.
"""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap, qn

import bookir as ir
import chunk as chunking
import merge as merging
import qa
from build_docx import Builder, add_arguments
from read_docx import read_docx

_W = f'xmlns:w="{nsmap["w"]}"'

#: What the fixture prints on the page, and the only place these strings are
#: written down. Every assertion derives from this table, so the fixture cannot
#: change under a test that still claims to describe it.
TITLE = "Pride and Prejudice"
VOLUME = "Volume One"
FRONTISPIECE = "Chapman and Hall"
PAGE_FIELD = " PAGE "
#: The digits Word had cached in the field when the file was saved. They are the
#: one thing in the header that must *not* come out the other end: a page number
#: is recomputed, never translated.
CACHED_PAGE = "137"
OPENERS = ["Prose that opens the first section.",
           "Prose that opens the second section."]

PERSIAN_TITLE = "غرور و تعصب"
PERSIAN_VOLUME = "دفتر یکم"
PERSIAN_FRONTISPIECE = "چاپمن و هال"


def _field(paragraph, instruction: str, cached: str) -> None:
    """A complex field, the shape Word writes a page number in."""
    for fragment in (
        '<w:fldChar w:fldCharType="begin"/>',
        f'<w:instrText xml:space="preserve">{instruction}</w:instrText>',
        '<w:fldChar w:fldCharType="separate"/>',
        f"<w:t>{cached}</w:t>",
        '<w:fldChar w:fldCharType="end"/>',
    ):
        paragraph._p.append(parse_xml(f"<w:r {_W}>{fragment}</w:r>"))


def _tab(paragraph) -> None:
    paragraph._p.append(parse_xml(f"<w:r {_W}><w:tab/></w:r>"))


@pytest.fixture(scope="module")
def headed_docx(tmp_path_factory) -> Path:
    """Two sections; the first defines the heads, the second inherits them."""
    document = Document()
    section = document.sections[0]
    section.different_first_page_header_footer = True

    section.header.is_linked_to_previous = False
    running = section.header.paragraphs[0]
    running.add_run(TITLE)
    _tab(running)
    _field(running, PAGE_FIELD, CACHED_PAGE)
    running._p.get_or_add_pPr().get_or_add_jc().set(qn("w:val"), "center")

    section.first_page_header.is_linked_to_previous = False
    section.first_page_header.paragraphs[0].add_run(FRONTISPIECE)

    section.footer.is_linked_to_previous = False
    section.footer.paragraphs[0].add_run(VOLUME)

    document.add_paragraph(OPENERS[0])
    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_paragraph(OPENERS[1])

    destination = tmp_path_factory.mktemp("running-heads") / "headed.docx"
    document.save(str(destination))
    return destination


@pytest.fixture(scope="module")
def imported(headed_docx, tmp_path_factory) -> dict:
    return read_docx(str(headed_docx), tmp_path_factory.mktemp("head-assets"))


def _translate(book: dict) -> dict:
    """Persian for every unit, so nothing is left for the builder to warn about."""
    for block in ir.iter_text_blocks(book):
        if (block.get("text") or "").strip():
            block["target"] = f"ترجمهٔ بند {block['id']} با طول کافی."
    persian = {TITLE: PERSIAN_TITLE, VOLUME: PERSIAN_VOLUME,
               FRONTISPIECE: PERSIAN_FRONTISPIECE}
    for _, _, piece, _ in ir.iter_running_pieces(book):
        piece["target"] = persian[ir.plain_text(piece["text"])]
    return book


def _build(book: dict, destination: Path) -> dict:
    parser = argparse.ArgumentParser()
    add_arguments(parser)
    options = parser.parse_args(["--book", "x", "--out", "y",
                                 "--font", "Tahoma", "--no-toc"])
    return Builder(book, destination.parent / "assets", options).build(destination)


def _running_parts(path: Path) -> dict[str, str]:
    """Every ``word/headerN.xml`` and ``word/footerN.xml`` of a built document."""
    with zipfile.ZipFile(path) as archive:
        return {name: archive.read(name).decode("utf-8")
                for name in archive.namelist()
                if re.fullmatch(r"word/(header|footer)\d+\.xml", name)}


def _outline(part: dict) -> list[str]:
    """What each piece of a running head is, in the order it is laid out."""
    return ["tab" if piece.get("tab") else "field" if piece.get("field") else "text"
            for line in part["paragraphs"] for piece in line["pieces"]]


# --------------------------------------------------------------------------- #
# Reading
# --------------------------------------------------------------------------- #

def test_the_authors_running_head_and_foot_reach_the_ir_as_prose(imported):
    """Both ends of the page, and both as text somebody can be asked to render."""
    carried = {(kind, ir.plain_text(piece["text"]))
               for _, kind, piece, _ in ir.iter_running_pieces(imported)}
    assert carried == {("header", TITLE), ("header", FRONTISPIECE),
                       ("footer", VOLUME)}


def test_a_page_number_is_carried_as_a_field_not_as_the_digits_it_showed(imported):
    """The digits were a cached result, not the author's words.

    Translating them would freeze the running head at whatever page the source
    was last saved on, and would put an English 137 in a Persian book besides.
    """
    head = imported["sections"][0]["headers"]["default"]
    assert _outline(head) == ["text", "tab", "field"]

    fields = [piece["field"] for line in head["paragraphs"]
              for piece in line["pieces"] if piece.get("field")]
    assert fields == [PAGE_FIELD]
    assert CACHED_PAGE not in json.dumps(imported["sections"], ensure_ascii=False)


def test_the_head_keeps_the_edge_the_author_put_it_on(imported):
    """``w:jc`` travels as its own token: ``center`` here, but ``start`` elsewhere."""
    head = imported["sections"][0]["headers"]["default"]
    assert [line["align"] for line in head["paragraphs"]] == ["center"]


def test_a_section_that_inherits_its_head_is_not_given_a_copy_of_it(imported):
    """Inheritance is how Word says it, and the IR says it the same way.

    Copying the head onto the second section would translate it twice, print it
    from two places, and let a later edit to one of them silently disagree with
    the other.
    """
    second = imported["sections"][1]
    assert (second["headers"], second["footers"]) == ({}, {})
    assert len(list(ir.iter_running_pieces(imported))) == 3


def test_a_first_page_head_the_section_never_shows_is_not_carried(tmp_path):
    """A definition Word ignores is not a running head; it is dead weight.

    Carrying it would put a head on the Persian edition's first page that the
    source never printed there.
    """
    document = Document()
    section = document.sections[0]
    section.first_page_header.is_linked_to_previous = False
    section.first_page_header.paragraphs[0].add_run(FRONTISPIECE)
    section.different_first_page_header_footer = False
    document.add_paragraph(OPENERS[0])
    destination = tmp_path / "unused-first.docx"
    document.save(str(destination))

    book = read_docx(str(destination), tmp_path / "assets")
    assert not list(ir.iter_running_pieces(book))
    assert book["sections"][0]["different_first_page"] is False


def test_a_picture_in_a_running_head_is_named_rather_than_quietly_lost(
        tmp_path, sample_png):
    """The words beside it are carried; the logo is not, and says so."""
    document = Document()
    header = document.sections[0].header
    header.is_linked_to_previous = False
    header.paragraphs[0].add_run(TITLE)
    header.paragraphs[0].add_run().add_picture(str(sample_png))
    document.add_paragraph(OPENERS[0])
    destination = tmp_path / "logo-head.docx"
    document.save(str(destination))

    book = read_docx(str(destination), tmp_path / "assets")
    warned = {w["kind"]: w for w in (book["source"].get("docx_warnings") or [])}
    assert "running-head-picture-dropped" in warned
    assert warned["running-head-picture-dropped"]["where"] \
        == "section 0 header (default)"
    assert [ir.plain_text(piece["text"])
            for _, _, piece, _ in ir.iter_running_pieces(book)] == [TITLE]


# --------------------------------------------------------------------------- #
# Out to a translator, and back
# --------------------------------------------------------------------------- #

def test_the_head_goes_out_as_a_worksheet_unit_and_merge_writes_it_back(
        imported, tmp_path):
    """The same ``@@ id kind`` protocol as every other string, not a second one."""
    book_path = tmp_path / "book.json"
    ir.save_book(dict(imported), book_path)
    chunks = tmp_path / "chunks"
    manifest = chunking.build(book_path, chunks, glossary_path=None, budget=100_000)

    expected = {unit_id for unit_id, _, _, _ in ir.iter_running_pieces(imported)}
    offered = {u for entry in manifest["chunks"] for u in entry["unit_ids"]}
    assert expected <= offered, "a running head nobody is asked to translate"

    worksheet = (chunks / manifest["chunks"][0]["file"]).read_text(encoding="utf-8")
    assert re.search(r"^@@ \S+ header$", worksheet, re.M)
    assert re.search(r"^@@ \S+ footer$", worksheet, re.M)

    for entry in manifest["chunks"]:
        (chunks / entry["output"]).write_text(
            "\n".join(f"@@ {unit_id} x\nترجمهٔ {unit_id}\n"
                      for unit_id in entry["unit_ids"]),
            encoding="utf-8", newline="",
        )
    report = merging.merge(book_path, chunks)
    assert report["ok"], report

    merged = ir.load_book(book_path)
    assert {piece["target"] for _, _, piece, _ in ir.iter_running_pieces(merged)} \
        == {f"ترجمهٔ {unit_id}" for unit_id in expected}


def test_a_running_head_nobody_translated_is_an_error_and_not_a_pass(imported):
    """Untranslated is the state the old decision existed to avoid printing."""
    report = qa.check_book(dict(imported)).summary()
    assert report["counts"]["running_heads"] == 3
    assert report["counts"]["running_heads_translated"] == 0
    assert report["by_code"]["untranslated-running-head"] == 3


def test_an_english_head_that_came_back_unchanged_is_caught_as_untranslated():
    """A filled-in target is not the same claim as a translated one.

    The typography pass reads every translated string in the book, and a running
    head is now one of them — so the English title answered back verbatim trips
    the gate an English paragraph would.
    """
    book = ir.new_book()
    book["sections"] = [{"index": 0, "start_block": None, "footers": {},
                         "headers": {"default": {"paragraphs": [
                             {"align": None,
                              "pieces": [{"id": "rh0001", "text": TITLE,
                                          "target": TITLE}]}]}}}]

    codes = {finding["code"]
             for finding in qa.check_book(book).summary()["findings"]
             if finding["unit"] == "rh0001"}
    assert "untranslated" in codes


# --------------------------------------------------------------------------- #
# Onto the page
# --------------------------------------------------------------------------- #

def test_the_persian_head_is_written_into_the_documents_own_header_part(
        imported, tmp_path):
    """Where a running head lives: ``word/headerN.xml``, not the body."""
    destination = tmp_path / "headed.fa.docx"
    report = _build(_translate(dict(imported)), destination)
    assert report["warning_count"] == 0, report["warnings"]

    parts = _running_parts(destination)
    running = next(body for body in parts.values() if PERSIAN_TITLE in body)
    assert "<w:bidi/>" in running and "<w:rtl/>" in running
    assert 'w:jc w:val="center"' in running
    assert f">{PAGE_FIELD}</w:instrText>" in running
    assert CACHED_PAGE not in running
    assert any(PERSIAN_VOLUME in body for body in parts.values())
    assert any(PERSIAN_FRONTISPIECE in body for body in parts.values())


def test_no_english_reaches_the_page_when_the_head_was_not_translated(
        imported, tmp_path):
    """The old decision, applied to the one case it still governs.

    A head with no Persian is left off the page and named in the report. Falling
    back to the source there instead would print the English title on every
    sheet, which is exactly what dropping them wholesale was avoiding.
    """
    book = _translate(dict(imported))
    head = next(piece for _, _, piece, _ in ir.iter_running_pieces(book)
                if ir.plain_text(piece["text"]) == TITLE)
    head["target"] = None

    destination = tmp_path / "half.fa.docx"
    report = _build(book, destination)

    assert not any(TITLE in body for body in _running_parts(destination).values())
    assert any(head["id"] in warning for warning in report["warnings"])


def test_the_generated_page_number_still_appears_when_the_source_has_no_foot(
        tmp_path):
    """The one behaviour that must not move: --page-numbers as it always was."""
    document = Document()
    document.sections[0].header.is_linked_to_previous = False
    document.sections[0].header.paragraphs[0].add_run(TITLE)
    document.add_paragraph(OPENERS[0])
    source = tmp_path / "head-only.docx"
    document.save(str(source))

    book = read_docx(str(source), tmp_path / "assets")
    for _, _, piece, _ in ir.iter_running_pieces(book):
        piece["target"] = PERSIAN_TITLE
    for block in ir.iter_text_blocks(book):
        block["target"] = "ترجمهٔ بندی با طول کافی برای آزمون."

    destination = tmp_path / "head-only.fa.docx"
    report = _build(book, destination)

    assert report["layout"]["page_numbers"] is True
    feet = [body for name, body in _running_parts(destination).items()
            if "footer" in name]
    assert feet and all("PAGE" in body for body in feet)


def test_a_source_foot_is_kept_instead_of_the_generated_page_number(
        imported, tmp_path):
    """Two things where the author put one is not fidelity.

    The generated number exists because the source had no foot of its own. This
    source does, and it may well be a page number already.
    """
    destination = tmp_path / "own-foot.fa.docx"
    report = _build(_translate(dict(imported)), destination)

    assert report["layout"]["page_numbers"] is False
    foot = next(body for name, body in _running_parts(destination).items()
                if "footer" in name)
    assert PERSIAN_VOLUME in foot
    assert "PAGE" not in foot


def test_the_running_heads_survive_the_round_trip(imported, tmp_path):
    """Read a Word file, build the Persian one, read that: same heads, same shape."""
    destination = tmp_path / "round.fa.docx"
    _build(_translate(dict(imported)), destination)

    rebuilt = read_docx(str(destination), tmp_path / "assets-back")
    assert _outline(rebuilt["sections"][0]["headers"]["default"]) \
        == _outline(imported["sections"][0]["headers"]["default"])
    assert {(kind, ir.plain_text(piece["text"]))
            for _, kind, piece, _ in ir.iter_running_pieces(rebuilt)} \
        == {("header", PERSIAN_TITLE), ("header", PERSIAN_FRONTISPIECE),
            ("footer", PERSIAN_VOLUME)}
    assert (rebuilt["sections"][1]["headers"],
            rebuilt["sections"][1]["footers"]) == ({}, {})
