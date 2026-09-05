"""The Word document: real footnotes, real bookmarks, real RTL, exact pictures.

These assertions read the saved package rather than the builder's own report,
because the failure that matters is a file Word cannot open or silently
renders wrong — not an in-memory object that looked right on the way out.
"""

from __future__ import annotations

import argparse
import zipfile
from pathlib import Path

import pytest

import bookir as ir
import qa
from build_docx import Builder, add_arguments, split_by_script

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _options(**overrides) -> argparse.Namespace:
    parser = argparse.ArgumentParser()
    add_arguments(parser)
    args = parser.parse_args(["--book", "x", "--out", "y", "--font", "Tahoma"])
    for key, value in overrides.items():
        setattr(args, key, value)
    return args


@pytest.fixture
def built(translated_book, tmp_path):
    book, assets = translated_book
    destination = tmp_path / "book.fa.docx"
    report = Builder(book, assets, _options()).build(destination)
    return book, destination, report


def test_build_reports_no_warnings(built):
    _book, _path, report = built
    assert report["warning_count"] == 0, report["warnings"]
    assert report["footnotes"] == 1
    assert report["bookmarks"] == report["headings"] >= 2


def test_package_has_a_real_footnotes_part(built):
    _book, path, _report = built
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        assert "word/footnotes.xml" in names
        content_types = archive.read("[Content_Types].xml").decode()
        assert "footnotes+xml" in content_types
        rels = archive.read("word/_rels/document.xml.rels").decode()
        assert "footnotes" in rels

        footnotes = archive.read("word/footnotes.xml").decode()
        # ids -1 and 0 are Word's separators; content starts at 1.
        assert 'w:id="-1"' in footnotes and 'w:id="0"' in footnotes
        assert 'w:id="1"' in footnotes
        assert "یادداشت مترجم" in footnotes
        assert "<w:footnoteRef/>" in footnotes

        document = archive.read("word/document.xml").decode()
        assert "<w:footnoteReference" in document


def test_document_is_right_to_left_at_every_level(built):
    _book, path, _report = built
    with zipfile.ZipFile(path) as archive:
        document = archive.read("word/document.xml").decode()
        styles = archive.read("word/styles.xml").decode()
    assert "<w:bidi/>" in document           # paragraphs
    assert "<w:rtl/>" in document            # Persian runs
    assert "<w:bidi/>" in styles             # Normal and heading styles
    assert 'w:cs="Tahoma"' in styles or 'w:cs="Tahoma"' in document


def test_headings_carry_bookmarks_that_the_toc_links_to(built):
    _book, path, _report = built
    with zipfile.ZipFile(path) as archive:
        document = archive.read("word/document.xml").decode()
    assert "<w:bookmarkStart" in document
    assert 'TOC \\o "1-' in document
    assert "w:anchor=" in document
    assert qa.check_docx(path).summary()["ok"]


def test_pictures_are_placed_with_an_explicit_extent(built):
    _book, path, _report = built
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        document = archive.read("word/document.xml").decode()
    assert any(n.startswith("word/media/") for n in names)
    assert "<wp:extent" in document


def test_document_reopens_with_persian_intact(built):
    from docx import Document

    _book, path, _report = built
    document = Document(str(path))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "کتابی کوچک" in text
    assert "متن فارسی" in text


def test_page_setup_comes_from_the_book(translated_book, tmp_path):
    from docx import Document

    book, assets = translated_book
    book["page"].update({"width_pt": 360.0, "height_pt": 600.0, "margin_top_pt": 40.0})
    destination = tmp_path / "sized.docx"
    Builder(book, assets, _options()).build(destination)

    section = Document(str(destination)).sections[0]
    assert section.page_width.pt == pytest.approx(360.0, abs=0.5)
    assert section.page_height.pt == pytest.approx(600.0, abs=0.5)
    assert section.top_margin.pt == pytest.approx(40.0, abs=0.5)


def test_untranslated_block_falls_back_to_source_and_warns(translated_book, tmp_path):
    book, assets = translated_book
    target_block = next(b for b in ir.iter_text_blocks(book) if b["type"] == "paragraph")
    target_block["target"] = None

    report = Builder(book, assets, _options()).build(tmp_path / "partial.docx")
    assert any("untranslated" in w for w in report["warnings"])


# --------------------------------------------------------------------------- #
# Bidi run splitting
# --------------------------------------------------------------------------- #

@pytest.mark.parametrize(
    "text, expected_latin",
    [
        ("الیزابت بنت (Elizabeth Bennet) رفت", ["Elizabeth Bennet"]),
        ("متن فارسی بدون لاتین", []),
        ("Alice در سرزمین Wonderland", ["Alice", "Wonderland"]),
    ],
)
def test_split_by_script_isolates_latin_runs(text, expected_latin):
    latin = [chunk.strip() for chunk, is_latin in split_by_script(text) if is_latin]
    assert latin == expected_latin
    assert "".join(chunk for chunk, _ in split_by_script(text)) == text


# --------------------------------------------------------------------------- #
# QA gates
# --------------------------------------------------------------------------- #

def test_qa_check_passes_a_complete_book(translated_book, tmp_path):
    book, assets = translated_book
    summary = qa.check_book(book, assets=assets).summary()
    assert summary["ok"], summary["findings"]


def test_qa_catches_a_dropped_footnote_marker(translated_book):
    book, assets = translated_book
    block = next(
        b for b in ir.iter_text_blocks(book) if ir.footnote_refs(b.get("text") or "")
    )
    block["target"] = "ترجمه‌ای که نشانهٔ پانویس را انداخته است."
    summary = qa.check_book(book, assets=assets).summary()
    assert not summary["ok"]
    assert "footnote-marker-lost" in summary["by_code"]


def test_qa_catches_an_omission_by_length_ratio(translated_book):
    book, assets = translated_book
    block = next(b for b in ir.iter_text_blocks(book) if b["type"] == "paragraph")
    block["text"] = "word " * 60
    block["target"] = "کوتاه."
    summary = qa.check_book(book, assets=assets).summary()
    assert "possible-omission" in summary["by_code"]


def test_qa_catches_a_tampered_image(translated_book, tmp_path):
    book, assets = translated_book
    image = next(b for b in book["blocks"] if b["type"] == "image")
    (assets / image["asset"]).write_bytes(b"not the original picture")
    summary = qa.check_book(book, assets=assets).summary()
    assert "asset-modified" in summary["by_code"]


def test_qa_docx_detects_a_link_with_no_bookmark(tmp_path):
    """A TOC entry pointing at nothing is exactly the bug users notice last."""
    source = tmp_path / "broken.docx"
    from docx import Document

    document = Document()
    paragraph = document.add_paragraph()
    import ooxml
    ooxml.add_internal_link(paragraph, "rv_9999", "فصل گمشده")
    document.save(str(source))

    summary = qa.check_docx(source).summary()
    assert not summary["ok"]
    assert "dead-link" in summary["by_code"]
