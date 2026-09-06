"""Section breaks and their page geometry, carried both ways.

A Word file that changes page size, turns a page on its side or re-cuts its
margins does so with a section break, and the reader used to keep only the
first section's geometry and say so in a ``sections-collapsed`` warning. A
warning is not preservation: a landscape plate came back portrait, an atlas
page came back the size of the prose around it, and the report said the file
had been read.

So these tests assert against the *source document*: this section was this
size and this way round, therefore the built book has a section that is too.
"""

from __future__ import annotations

import argparse
from pathlib import Path

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.oxml.ns import qn
from docx.shared import Pt

import bookir as ir
from build_docx import Builder, add_arguments
from read_docx import read_docx

#: The three shapes the fixture cuts, in order. Every assertion derives from
#: this table rather than repeating its numbers, so a change here cannot leave
#: a test asserting geometry the file no longer has.
SHAPES = [
    {"start_type": "nextPage", "orientation": "portrait",
     "width_pt": 420.0, "height_pt": 640.0,
     "margin_top_pt": 50.0, "margin_bottom_pt": 55.0,
     "margin_inner_pt": 60.0, "margin_outer_pt": 45.0},
    {"start_type": "nextPage", "orientation": "landscape",
     "width_pt": 720.0, "height_pt": 500.0,
     "margin_top_pt": 20.0, "margin_bottom_pt": 21.0,
     "margin_inner_pt": 22.0, "margin_outer_pt": 23.0},
    {"start_type": "continuous", "orientation": "portrait",
     "width_pt": 380.0, "height_pt": 700.0,
     "margin_top_pt": 30.0, "margin_bottom_pt": 31.0,
     "margin_inner_pt": 32.0, "margin_outer_pt": 33.0},
]

#: What the builder reproduces from a section record. Gutter and the header and
#: footer distances are deliberately not in it: those are house style, set from
#: the command line for every section alike, and the IR keeps the source's own
#: values whether or not the build uses them.
REBUILT = ("start_type", "orientation", "width_pt", "height_pt",
           "margin_top_pt", "margin_bottom_pt", "margin_inner_pt",
           "margin_outer_pt")

OPENERS = ["Prose in the first section.",
           "Prose in the second section.",
           "Prose in the third section."]


def _shape(section, values: dict) -> None:
    section.page_width = Pt(values["width_pt"])
    section.page_height = Pt(values["height_pt"])
    section.top_margin = Pt(values["margin_top_pt"])
    section.bottom_margin = Pt(values["margin_bottom_pt"])
    section.left_margin = Pt(values["margin_inner_pt"])
    section.right_margin = Pt(values["margin_outer_pt"])
    section.orientation = (WD_ORIENT.LANDSCAPE
                           if values["orientation"] == "landscape"
                           else WD_ORIENT.PORTRAIT)


@pytest.fixture(scope="module")
def sectioned_docx(tmp_path_factory) -> Path:
    """Three real sections, no two of them the same shape."""
    document = Document()
    document.sections[0].gutter = Pt(18)
    document.sections[0].header_distance = Pt(30)
    document.sections[0].footer_distance = Pt(24)

    for number, values in enumerate(SHAPES):
        section = (document.sections[0] if number == 0
                   else document.add_section(
                       WD_SECTION_START.from_xml(values["start_type"])))
        _shape(section, values)
        document.add_paragraph(OPENERS[number])
        document.add_paragraph(f"More prose still inside section {number}.")

    destination = tmp_path_factory.mktemp("docx-sections") / "sectioned.docx"
    document.save(str(destination))
    return destination


@pytest.fixture(scope="module")
def imported(sectioned_docx, tmp_path_factory):
    return read_docx(str(sectioned_docx), tmp_path_factory.mktemp("section-assets"))


def _translate(book: dict) -> dict:
    """A stub translation, so no block is left for the builder to warn about."""
    for block in ir.iter_text_blocks(book):
        if (block.get("text") or "").strip():
            block["target"] = f"ترجمهٔ بند {block['id']} با طول کافی."
    return book


def _build(book, assets, destination: Path) -> dict:
    parser = argparse.ArgumentParser()
    add_arguments(parser)
    options = parser.parse_args(["--book", "x", "--out", "y",
                                 "--font", "Tahoma", "--no-toc"])
    return Builder(book, assets, options).build(destination)


# --------------------------------------------------------------------------- #
# Reading
# --------------------------------------------------------------------------- #

def test_every_section_reaches_the_ir_in_order(imported):
    assert [record["index"] for record in imported["sections"]] == list(
        range(len(SHAPES)))


def test_each_section_keeps_its_own_geometry(imported):
    """One shape for the whole file was the loss; three shapes is the fix."""
    carried = [{field: record[field] for field in REBUILT}
               for record in imported["sections"]]
    assert carried == SHAPES


def test_a_section_names_the_block_it_opens_at(imported):
    """Off by one here and the break lands inside the previous section's prose."""
    blocks = {block["id"]: ir.plain_text(block["text"] or "")
              for block in imported["blocks"]}
    opened = [blocks[record["start_block"]] for record in imported["sections"]]
    assert opened == OPENERS


def test_the_page_setup_is_still_the_first_section(imported):
    """Everything downstream reads `book["page"]`; it may not have moved.

    Derived from the section record rather than restated, so the two cannot
    drift apart without this failing.
    """
    first = imported["sections"][0]
    assert imported["page"] == {field: first[field] for field in imported["page"]}


def test_the_binding_measurements_are_carried_too(imported):
    """Not used by the builder, which sets house style — but not thrown away."""
    first = imported["sections"][0]
    assert (first["gutter_pt"], first["header_distance_pt"],
            first["footer_distance_pt"]) == (18.0, 30.0, 24.0)


def test_the_collapse_warning_is_not_told_any_more(imported):
    """It described a real loss once. Repeating it now would be a lie."""
    kinds = {w["kind"] for w in (imported["source"].get("docx_warnings") or [])}
    assert "sections-collapsed" not in kinds


def test_a_break_folded_into_a_content_paragraph_starts_the_next_block(tmp_path):
    """Word gives the break a paragraph of its own; converters do not.

    The ``w:sectPr`` then sits in the ``w:pPr`` of the last paragraph *of* the
    section, which still belongs to it — the next section opens at the block
    after that one, not at it.
    """
    document = Document()
    document.add_paragraph("Last paragraph of the first section.")
    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_paragraph("First paragraph of the second section.")

    body = document.element.body
    carrier = next(p for p in body.findall(qn("w:p"))
                   if p.find(qn("w:pPr")) is not None
                   and p.find(qn("w:pPr")).find(qn("w:sectPr")) is not None)
    properties = carrier.find(qn("w:pPr"))
    carrier.getprevious().get_or_add_pPr().append(properties.find(qn("w:sectPr")))
    body.remove(carrier)

    destination = tmp_path / "folded.docx"
    document.save(str(destination))

    book = read_docx(str(destination), tmp_path / "assets")
    opened = book["sections"][1]["start_block"]
    text = next(b["text"] for b in book["blocks"] if b["id"] == opened)
    assert ir.plain_text(text) == "First paragraph of the second section."


def test_a_single_section_document_reports_exactly_one(tmp_path):
    document = Document()
    document.add_paragraph("Nothing but one section of prose.")
    destination = tmp_path / "plain.docx"
    document.save(str(destination))

    book = read_docx(str(destination), tmp_path / "assets")
    assert len(book["sections"]) == 1
    assert book["sections"][0]["start_block"] == book["blocks"][0]["id"]


def test_columns_are_the_one_thing_still_named_as_dropped(tmp_path):
    """The general lie is gone; what is genuinely not carried is still said."""
    document = Document()
    document.add_paragraph("Two columns of prose.")
    document.sections[0]._sectPr.find(qn("w:cols")).set(qn("w:num"), "2")
    destination = tmp_path / "columns.docx"
    document.save(str(destination))

    book = read_docx(str(destination), tmp_path / "assets")
    warned = {w["kind"]: w for w in (book["source"].get("docx_warnings") or [])}
    assert "section-columns-dropped" in warned
    assert warned["section-columns-dropped"]["count"] == 2


# --------------------------------------------------------------------------- #
# And back out again
# --------------------------------------------------------------------------- #

def test_the_sections_survive_the_round_trip(imported, tmp_path):
    """Read a Word file, build the Persian one, read that: the shapes match."""
    book = _translate(dict(imported))
    destination = tmp_path / "sectioned.fa.docx"
    report = _build(book, tmp_path / "assets", destination)
    assert report["warning_count"] == 0, report["warnings"]

    rebuilt = read_docx(str(destination), tmp_path / "assets-back")
    carried = [{field: record[field] for field in REBUILT}
               for record in rebuilt["sections"]]
    assert carried == SHAPES


def test_the_break_lands_in_front_of_the_block_that_opens_the_section(
        imported, tmp_path):
    """A section with the right geometry in the wrong place is still wrong."""
    book = _translate(dict(imported))
    destination = tmp_path / "placed.fa.docx"
    _build(book, tmp_path / "assets", destination)

    expected = [next(b["target"] for b in book["blocks"]
                     if b["id"] == record["start_block"])
                for record in book["sections"][1:]]

    rebuilt = read_docx(str(destination), tmp_path / "assets-back")
    blocks = {b["id"]: b.get("text") or "" for b in rebuilt["blocks"]}
    opened = [blocks[record["start_block"]] for record in rebuilt["sections"][1:]]
    assert opened == expected


def test_a_book_with_no_sections_still_builds_one(tmp_path):
    """Every other reader produces a book with no `sections` key at all."""
    book = ir.new_book()
    block = ir.make_block("paragraph", 1, page=0, text="Plain prose.")
    block["target"] = "متن ساده‌ای که هیچ بخش‌بندی‌ای ندارد."
    book["blocks"] = [block]

    destination = tmp_path / "sectionless.docx"
    _build(book, tmp_path, destination)
    assert len(Document(destination).sections) == 1
