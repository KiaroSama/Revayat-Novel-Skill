"""Hyperlinks rebuilt as hyperlinks, and named when they cannot be.

The reader keeps a link's words in the prose and its target on the block. The
builder used to write only the words, so an edition's citations arrived as
sentences that mention a source nobody can follow.

Putting them back is limited by translation itself: the translator is never
shown a URL, so the only thing that can locate a link in a Persian sentence is
the display phrase surviving word for word — a name, a title, a number, an
address written out as text. Ordinary prose does not survive that way, and is
not meant to. So the tests here come in two halves: the links that can be put
back must be real Word hyperlinks pointing at the original target, and the ones
that cannot must be named individually rather than counted.
"""

from __future__ import annotations

import argparse
import re
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.hyperlink import Hyperlink

import bookir as ir
import qa
from build_docx import Builder, add_arguments
from read_docx import read_docx

EXTERNAL = "https://example.com/a/source?q=1#part"
ANCHOR = "#the_target_bookmark"

#: A name a translator would carry across untouched, which is exactly the case
#: this feature can rescue.
KEPT = "Elizabeth Bennet"


def _source_link(paragraph, text: str, href: str) -> None:
    """A ``w:hyperlink`` in a *source* document, external or in-document.

    python-docx reads hyperlinks but cannot create one, and the fixture has to
    contain the shape the reader is being tested on.
    """
    link = OxmlElement("w:hyperlink")
    if href.startswith("#"):
        link.set(qn("w:anchor"), href[1:])
    else:
        link.set(qn("r:id"),
                 paragraph.part.relate_to(href, RT.HYPERLINK, is_external=True))
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _source_bookmark(paragraph, name: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "77")
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "77")
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _book(target: str, links: list[dict[str, str]], *,
          bookmarks: list[str] | None = None) -> dict:
    """One translated paragraph carrying link metadata, plus its landing block."""
    book = ir.new_book()
    block = ir.make_block("paragraph", 1, page=0, text="Source prose.")
    block["target"] = target
    block["links"] = links

    landing = ir.make_block("paragraph", 2, page=0, text="The destination.")
    landing["target"] = "بندی که پیوند درون‌متنی به آن اشاره می‌کند."
    if bookmarks:
        landing["bookmarks"] = bookmarks

    book["blocks"] = [block, landing]
    return book


def _build(book: dict, destination: Path) -> dict:
    parser = argparse.ArgumentParser()
    add_arguments(parser)
    options = parser.parse_args(["--book", "x", "--out", "y",
                                 "--font", "Tahoma", "--no-toc"])
    return Builder(book, destination.parent, options).build(destination)


def _links_in(path: Path) -> list[dict[str, str]]:
    """Every link the built document holds, read back by python-docx itself."""
    found = []
    for paragraph in Document(path).paragraphs:
        for item in paragraph.iter_inner_content():
            if isinstance(item, Hyperlink):
                found.append({"text": item.text,
                              "href": (f"#{item.fragment}" if not item.address
                                       else item.address)})
    return found


def _document_xml(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        return archive.read("word/document.xml").decode("utf-8")


# --------------------------------------------------------------------------- #
# Reading: the target and the place it points at
# --------------------------------------------------------------------------- #

def test_an_in_document_link_keeps_its_anchor_and_its_destination(tmp_path):
    """An anchor whose bookmark was dropped is a dead link, not a preserved one."""
    document = Document()
    paragraph = document.add_paragraph("See ")
    _source_link(paragraph, "Table 7", ANCHOR)
    destination = document.add_paragraph("")
    _source_bookmark(destination, ANCHOR[1:])
    destination.add_run("The table itself.")

    path = tmp_path / "anchored.docx"
    document.save(str(path))

    book = read_docx(str(path), tmp_path / "assets")
    assert book["blocks"][0]["links"] == [{"text": "Table 7", "href": ANCHOR}]
    assert book["blocks"][1]["bookmarks"] == [ANCHOR[1:]]


def test_bookmarks_nothing_links_to_are_left_behind(tmp_path):
    """Word files are full of `_GoBack` and `_Toc…`; none of them is a link."""
    document = Document()
    paragraph = document.add_paragraph("Prose with a stray anchor in it.")
    _source_bookmark(paragraph, "_GoBack")

    path = tmp_path / "stray.docx"
    document.save(str(path))

    book = read_docx(str(path), tmp_path / "assets")
    assert not any(b.get("bookmarks") for b in book["blocks"])


# --------------------------------------------------------------------------- #
# Building: what can be put back
# --------------------------------------------------------------------------- #

def test_an_external_link_is_rebuilt_with_its_target_intact(tmp_path):
    book = _book(f"او دربارهٔ {KEPT} نوشت.", [{"text": KEPT, "href": EXTERNAL}])
    destination = tmp_path / "external.docx"
    report = _build(book, destination)

    assert report["warning_count"] == 0, report["warnings"]
    assert _links_in(destination) == [{"text": KEPT, "href": EXTERNAL}]


def test_the_external_relationship_is_marked_external(tmp_path):
    """An internal relationship to a URL makes Word refuse to open the file."""
    book = _book(f"او دربارهٔ {KEPT} نوشت.", [{"text": KEPT, "href": EXTERNAL}])
    destination = tmp_path / "external-rel.docx"
    _build(book, destination)

    with zipfile.ZipFile(destination) as archive:
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
    hyperlinks = [tag for tag in re.findall(r"<Relationship\b[^>]*>", rels)
                  if "/hyperlink" in tag]
    assert len(hyperlinks) == 1, hyperlinks
    assert 'TargetMode="External"' in hyperlinks[0]
    assert f'Target="{EXTERNAL}"'.replace("&", "&amp;") in hyperlinks[0]


def test_an_internal_anchor_is_rebuilt_and_has_somewhere_to_land(tmp_path):
    """The anchor alone is half of it: without the bookmark the link is dead."""
    book = _book("جدول ۷ را ببینید.", [{"text": "جدول ۷", "href": ANCHOR}],
                 bookmarks=[ANCHOR[1:]])
    destination = tmp_path / "anchor.docx"
    report = _build(book, destination)

    assert report["warning_count"] == 0, report["warnings"]
    assert _links_in(destination) == [{"text": "جدول ۷", "href": ANCHOR}]

    package = qa.check_docx(destination, book).summary()
    assert package["ok"], package["findings"]


def test_a_persian_phrase_inside_a_link_keeps_the_direction_of_the_prose(tmp_path):
    """A run in a `w:hyperlink` needs `w:rtl` exactly like the words around it."""
    book = _book("جدول ۷ را ببینید.", [{"text": "جدول ۷", "href": ANCHOR}],
                 bookmarks=[ANCHOR[1:]])
    destination = tmp_path / "rtl-link.docx"
    _build(book, destination)

    link = re.search(r"<w:hyperlink.*?</w:hyperlink>", _document_xml(destination),
                     re.S)
    assert link, "no hyperlink was written at all"
    assert "<w:rtl/>" in link.group(0), (
        "the linked words are left-to-right inside a right-to-left sentence"
    )


def test_a_latin_phrase_inside_a_link_stays_left_to_right(tmp_path):
    """The same rule the other way: a Latin name is not marked `w:rtl`."""
    book = _book(f"او دربارهٔ {KEPT} نوشت.", [{"text": KEPT, "href": EXTERNAL}])
    destination = tmp_path / "ltr-link.docx"
    _build(book, destination)

    link = re.search(r"<w:hyperlink.*?</w:hyperlink>", _document_xml(destination),
                     re.S)
    assert link and "<w:rtl/>" not in link.group(0)


# --------------------------------------------------------------------------- #
# And what cannot: every miss named, none of them guessed at
# --------------------------------------------------------------------------- #

def _refusal(report: dict, link: dict[str, str]) -> str:
    """The one warning about this link, or a failure saying it was not made."""
    named = [w for w in report["warnings"]
             if link["text"] in w and link["href"] in w]
    assert len(named) == 1, (
        f"expected exactly one warning naming {link!r}: {report['warnings']}"
    )
    return named[0]


def test_a_link_whose_words_did_not_survive_translation_is_named(tmp_path):
    """The ordinary case, and the reason the hit rate is low by design."""
    link = {"text": "the linked phrase", "href": EXTERNAL}
    book = _book("بندی که هیچ‌کدام از واژه‌های انگلیسی در آن نمانده است.", [link])
    destination = tmp_path / "lost.docx"
    report = _build(book, destination)

    _refusal(report, link)
    assert not _links_in(destination), "a link was placed on the wrong words"


def test_an_anchor_with_no_destination_in_the_book_is_refused(tmp_path):
    """Emitting it anyway would put a dead link past the package gate."""
    link = {"text": "جدول ۷", "href": "#never_carried"}
    book = _book("جدول ۷ را ببینید.", [link])
    destination = tmp_path / "dangling.docx"
    report = _build(book, destination)

    _refusal(report, link)
    assert not _links_in(destination)
    assert qa.check_docx(destination, book).summary()["ok"]


def test_two_links_sharing_the_same_words_are_both_refused(tmp_path):
    """Nothing in the sentence says which of the two targets belongs to it."""
    here = {"text": "here", "href": "https://example.com/one"}
    there = {"text": "here", "href": "https://example.com/two"}
    book = _book("بند فارسی با واژهٔ here در میان آن.", [here, there])
    destination = tmp_path / "shared.docx"
    report = _build(book, destination)

    _refusal(report, here)
    _refusal(report, there)
    assert not _links_in(destination)


def test_a_phrase_that_appears_twice_is_not_guessed_at(tmp_path):
    """Linking the first of two is a coin toss, and a wrong link reads as right."""
    link = {"text": KEPT, "href": EXTERNAL}
    book = _book(f"{KEPT} آمد و سپس {KEPT} رفت.", [link])
    destination = tmp_path / "ambiguous.docx"
    report = _build(book, destination)

    _refusal(report, link)
    assert not _links_in(destination)


def test_a_phrase_split_across_emphasis_is_named_rather_than_dropped(tmp_path):
    """It survived the translation, but not as one stretch of text to link."""
    link = {"text": KEPT, "href": EXTERNAL}
    book = _book("او دربارهٔ Elizabeth **Bennet** نوشت.", [link])
    destination = tmp_path / "split.docx"
    report = _build(book, destination)

    _refusal(report, link)
    assert not _links_in(destination)


def test_the_untranslated_source_still_gets_its_links(tmp_path):
    """A block the translator has not reached is written as-is; so is its link."""
    book = _book("", [{"text": KEPT, "href": EXTERNAL}])
    book["blocks"][0]["text"] = f"A sentence about {KEPT} and its source."
    book["blocks"][0]["target"] = None

    destination = tmp_path / "untranslated.docx"
    report = _build(book, destination)

    assert _links_in(destination) == [{"text": KEPT, "href": EXTERNAL}]
    assert any("untranslated" in warning for warning in report["warnings"])
