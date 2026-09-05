"""DOCX as a real input, carried end to end through the same pipeline.

The failures this file exists for are all *silent* ones. A DOCX that loses its
tables still imports, still translates, still builds, and still passes every
structural gate — because the gates compare the IR against itself, and the
words were gone before the IR existed. Nothing downstream can notice content
that was never read.

So these tests assert against the source document rather than the IR: this
sentence was in the Word file, therefore it must be in the book.

The fixture is generated, not committed. It is built once per session with
every structure §14 asks for — headings, prose, bold and italic, a hyperlink,
a nested list, an image, a table, a page break and a footnote — and the whole
pipeline is then run over it.
"""

from __future__ import annotations

import argparse
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

import bookir as ir
import qa
from read_docx import read_docx
from tests_support import png_bytes

LINKED = "the linked phrase"
IN_TABLE = "This cell text must reach the translation."
AFTER_BREAK = "This paragraph follows a hard page break."
FOOTNOTED = "A sentence carrying a note."


def _hyperlink(paragraph, text: str, url: str) -> None:
    """python-docx reads hyperlinks but cannot create one, so build the element.

    Worth doing rather than skipping: a linked phrase is exactly the shape that
    used to vanish, because `paragraph.runs` does not descend into `w:hyperlink`.
    """
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def _footnote(document, paragraph, text: str) -> None:
    """Attach a real ``word/footnotes.xml`` note and reference it."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    body = (
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/'
        '2006/main"><w:footnote w:id="1"><w:p><w:r><w:t>{text}</w:t></w:r></w:p>'
        "</w:footnote></w:footnotes>"
    ).format(text=text)
    part = Part(
        PackURI("/word/footnotes.xml"),
        "application/vnd.openxmlformats-officedocument.wordprocessingml."
        "footnotes+xml",
        body.encode("utf-8"),
        document.part.package,
    )
    document.part.relate_to(part, RT.FOOTNOTES)

    run = paragraph.add_run()
    reference = OxmlElement("w:footnoteReference")
    reference.set(qn("w:id"), "1")
    run._r.append(reference)


@pytest.fixture(scope="module")
def rich_docx(tmp_path_factory) -> Path:
    """One Word file with every structure the skill claims to read."""
    directory = tmp_path_factory.mktemp("docx-input")
    picture = directory / "plate.png"
    picture.write_bytes(png_bytes(160, 100))

    document = Document()
    document.add_heading("Chapter One", level=1)
    document.add_heading("A Section Within It", level=2)

    paragraph = document.add_paragraph("Ordinary prose with ")
    paragraph.add_run("bold").bold = True
    paragraph.add_run(" and ")
    paragraph.add_run("italic").italic = True
    paragraph.add_run(" in it, plus ")
    _hyperlink(paragraph, LINKED, "https://example.com/source")
    paragraph.add_run(" carried on afterwards.")

    noted = document.add_paragraph(FOOTNOTED)
    _footnote(document, noted, "The note the book itself carries.")

    document.add_paragraph("Top-level item", style="List Bullet")
    document.add_paragraph("Nested item", style="List Bullet 2")

    document.add_picture(str(picture), width=Pt(160))
    document.add_paragraph("The picture above.", style="Caption")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Column one"
    table.cell(0, 1).text = "Column two"
    table.cell(1, 0).text = IN_TABLE
    table.cell(1, 1).text = "A shorter cell."

    document.add_page_break()
    document.add_paragraph(AFTER_BREAK)

    destination = directory / "rich.docx"
    document.save(destination)
    return destination


@pytest.fixture(scope="module")
def imported(rich_docx, tmp_path_factory):
    assets = tmp_path_factory.mktemp("docx-assets")
    return read_docx(str(rich_docx), assets), assets


def _all_text(book) -> str:
    return "\n".join(b.get("text") or "" for b in book["blocks"])


# --------------------------------------------------------------------------- #
# Nothing may be lost on the way in
# --------------------------------------------------------------------------- #

def test_the_import_is_a_valid_book(imported):
    book, _ = imported
    assert ir.validate_book(book) == []
    assert book["source"]["format"] == "docx"


def test_table_cells_survive_and_know_where_they_came_from(imported):
    """They used to vanish: `document.paragraphs` never descends into a table."""
    book, _ = imported
    cells = [b for b in book["blocks"] if b.get("table")]
    assert IN_TABLE in _all_text(book), "the table's text was dropped"
    assert len(cells) == 4
    assert {(b["row"], b["cell"]) for b in cells} == {(1, 1), (1, 2), (2, 1), (2, 2)}


def test_a_flattened_table_is_declared_not_implied(imported):
    """The words are kept; the grid is not. Both facts have to be stated."""
    book, _ = imported
    warnings = book["source"].get("docx_warnings") or []
    flattened = [w for w in warnings if w["kind"] == "table-flattened"]
    assert flattened and flattened[0]["rows"] == 2


def test_linked_text_survives(imported):
    """`paragraph.runs` skips hyperlink content, so this phrase used to go."""
    book, _ = imported
    assert LINKED in _all_text(book)


def test_emphasis_survives_as_markup(imported):
    book, _ = imported
    prose = next(b for b in book["blocks"] if "Ordinary prose" in (b["text"] or ""))
    assert "**bold**" in prose["text"] and "*italic*" in prose["text"]
    assert ir.emphasis_signature(prose["text"])[:2] == (1, 1)


def test_heading_levels_come_through(imported):
    book, _ = imported
    headings = [(b["level"], ir.plain_text(b["text"]))
                for b in book["blocks"] if b["type"] == "heading"]
    assert headings == [(1, "Chapter One"), (2, "A Section Within It")]


def test_list_nesting_depth_is_read_from_the_numbering(imported):
    """The style says it is a list; only w:ilvl says how deep."""
    book, _ = imported
    items = [b["level"] for b in book["blocks"] if b["type"] == "listitem"]
    assert items == [1, 2], f"nesting was flattened: {items}"


def test_the_page_break_is_kept_and_the_order_holds(imported):
    book, _ = imported
    kinds = [b["type"] for b in book["blocks"]]
    assert "pagebreak" in kinds
    after = next(i for i, b in enumerate(book["blocks"])
                 if AFTER_BREAK in (b.get("text") or ""))
    assert kinds.index("pagebreak") < after


def test_the_image_is_extracted_with_its_bytes_and_size(imported):
    book, assets = imported
    image = next(b for b in book["blocks"] if b["type"] == "image")
    written = assets / image["asset"]
    assert written.exists()
    assert ir.sha256_bytes(written.read_bytes()) == image["sha256"]
    assert image["width_pt"] == pytest.approx(160, abs=1)


def test_the_caption_keeps_its_own_kind(imported):
    book, _ = imported
    assert any(b["type"] == "caption" for b in book["blocks"])


def test_the_footnote_is_read_and_anchored(imported):
    book, _ = imported
    assert len(book["footnotes"]) == 1
    note = book["footnotes"][0]
    assert note["origin"] == "source" and note["anchor_block"]
    anchor = next(b for b in book["blocks"] if b["id"] == note["anchor_block"])
    assert note["id"] in ir.footnote_refs(anchor["text"])


def test_page_geometry_comes_from_the_document(imported):
    book, _ = imported
    assert book["page"]["width_pt"] > 0 and book["page"]["height_pt"] > 0


# --------------------------------------------------------------------------- #
# And it has to go all the way through
# --------------------------------------------------------------------------- #

def test_a_docx_goes_through_the_whole_pipeline(imported, tmp_path):
    """Import, translate, gate, build, verify — the same route a PDF takes."""
    from build_docx import Builder, add_arguments

    book, assets = imported

    for block in ir.iter_text_blocks(book):
        if not (block.get("text") or "").strip():
            continue
        # A translator must carry every footnote marker across; a stub that
        # drops them would be testing a book no real run could produce.
        markers = "".join(f"[[fn:{note}]]"
                          for note in ir.footnote_refs(block["text"]))
        block["target"] = (
            "ترجمهٔ این بند که به اندازهٔ کافی بلند است تا از نسبت طول رد شود "
            f"({block['id']}){markers}."
        )
    for note in book["footnotes"]:
        note["target"] = "یادداشت ترجمه‌شده."
    book["meta"]["title_target"] = "فصل یکم"

    gate = qa.check_book(book, assets=assets).summary()
    assert gate["ok"], gate["findings"]

    parser = argparse.ArgumentParser()
    add_arguments(parser)
    options = parser.parse_args(["--book", "x", "--out", "y", "--font", "Tahoma"])
    destination = tmp_path / "from-docx.fa.docx"
    report = Builder(book, assets, options).build(destination)

    assert report["warning_count"] == 0, report["warnings"]
    assert report["footnotes"] == 1
    assert report["headings"] == 2

    package = qa.check_docx(destination, book).summary()
    assert package["ok"], package["findings"]

    with zipfile.ZipFile(destination) as archive:
        document = archive.read("word/document.xml").decode("utf-8")
    assert "<w:bidi" in document, "the Persian document is not right-to-left"


# --------------------------------------------------------------------------- #
# And the grid comes back
# --------------------------------------------------------------------------- #

def test_the_table_is_rebuilt_as_a_table_not_as_loose_paragraphs(imported,
                                                                 tmp_path):
    """The cells travel as blocks so each one is translated; the grid returns here.

    Keeping them as ordinary text blocks through the pipeline is what makes
    every cell its own worksheet unit — translated, counted and gated like a
    paragraph. Only the builder has anything to put the grid back into.
    """
    from docx import Document

    from build_docx import Builder, add_arguments

    book, assets = imported
    for block in ir.iter_text_blocks(book):
        if (block.get("text") or "").strip():
            markers = "".join(f"[[fn:{n}]]" for n in ir.footnote_refs(block["text"]))
            block["target"] = f"ترجمهٔ بند {block['id']} با طول کافی{markers}."
    for note in book["footnotes"]:
        note["target"] = "یادداشت."

    parser = argparse.ArgumentParser()
    add_arguments(parser)
    options = parser.parse_args(["--book", "x", "--out", "y", "--font", "Tahoma",
                                 "--no-toc"])
    destination = tmp_path / "with-table.docx"
    Builder(book, assets, options).build(destination)

    document = Document(destination)
    assert len(document.tables) == 1, "the grid was not rebuilt"
    table = document.tables[0]
    assert (len(table.rows), len(table.columns)) == (2, 2)

    # All four cells of the fixture carry text, and each one landed in its own
    # cell rather than being concatenated into a neighbour.
    filled = [table.cell(r, c).text.strip()
              for r in range(2) for c in range(2)]
    assert all(filled), filled
    assert len(set(filled)) == 4, f"cells were merged or duplicated: {filled}"


def test_a_persian_table_reads_right_to_left(imported, tmp_path):
    """Cell paragraphs alone do not do this — the column order is separate."""
    import zipfile

    from build_docx import Builder, add_arguments

    book, assets = imported
    for block in ir.iter_text_blocks(book):
        if (block.get("text") or "").strip():
            markers = "".join(f"[[fn:{n}]]" for n in ir.footnote_refs(block["text"]))
            block["target"] = f"ترجمهٔ بند {block['id']} با طول کافی{markers}."
    for note in book["footnotes"]:
        note["target"] = "یادداشت."

    parser = argparse.ArgumentParser()
    add_arguments(parser)
    options = parser.parse_args(["--book", "x", "--out", "y", "--font", "Tahoma",
                                 "--no-toc"])
    destination = tmp_path / "rtl-table.docx"
    Builder(book, assets, options).build(destination)

    with zipfile.ZipFile(destination) as archive:
        document = archive.read("word/document.xml").decode("utf-8")
    assert "<w:bidiVisual" in document, (
        "the table keeps left-to-right column order in a Persian document"
    )


def test_a_cell_with_no_coordinates_is_written_rather_than_dropped(tmp_path):
    """Losing a sentence to a malformed row is worse than an untidy table."""
    from build_docx import Builder, add_arguments

    book = ir.new_book()
    stray = ir.make_block("paragraph", 1, page=0, text="Orphaned cell text.",
                          table="t0001")
    stray["target"] = "متن سلولی که مختصاتش را از دست داده و باید بماند."
    book["blocks"] = [stray]

    parser = argparse.ArgumentParser()
    add_arguments(parser)
    options = parser.parse_args(["--book", "x", "--out", "y", "--font", "Tahoma",
                                 "--no-toc"])
    destination = tmp_path / "stray.docx"
    Builder(book, tmp_path, options).build(destination)

    from docx import Document
    document = Document(destination)
    assert not document.tables
    assert any("سلولی" in p.text for p in document.paragraphs)
